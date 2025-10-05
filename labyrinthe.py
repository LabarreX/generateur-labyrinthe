from random import *
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from reportlab.lib.pagesizes import landscape, A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
import io

class Labyrinthe: # Classe représentant un labyrinthe
    def __init__(self, largeur, hauteur): # Initialise un labyrinthe fermé
        self.largeur = largeur
        self.hauteur = hauteur
        self.cells = [[{'id': y * largeur + x, 'nord': False, 'sud': False, 'est': False, 'ouest': False} for x in range(largeur)] for y in range(hauteur)]
    
    def ouvrir(self, x1, y1, x2, y2): # Ouvre un mur entre deux cellules adjacentes
        if x1 == x2 and y1 == y2 + 1:
            self.cells[y1][x1]['nord'] = True
            self.cells[y2][x2]['sud'] = True
        elif x1 == x2 and y1 == y2 - 1:
            self.cells[y1][x1]['sud'] = True
            self.cells[y2][x2]['nord'] = True
        elif x1 == x2 + 1 and y1 == y2:
            self.cells[y1][x1]['ouest'] = True
            self.cells[y2][x2]['est'] = True
        elif x1 == x2 - 1 and y1 == y2:
            self.cells[y1][x1]['est'] = True
            self.cells[y2][x2]['ouest'] = True
        else:
            raise ValueError("Les cellules ne sont pas adjacentes")
    
    def generer_fusion(self): # Algorithme de fusion de cellules
        ids = [cell['id'] for row in self.cells for cell in row]
        while True: # Répéter jusqu'à ce que toutes les cellules aient le même id
            x = randint(0, self.largeur - 1) # Choisir une cellule au hasard
            y = randint(0, self.hauteur - 1)
            directions = []
            if x > 0:
                directions.append((-1, 0))
            if x < self.largeur - 1:
                directions.append((1, 0))
            if y > 0:
                directions.append((0, -1))
            if y < self.hauteur - 1:
                directions.append((0, 1))
            if not directions:
                continue
            dx, dy = choice(directions) # Choisir une direction au hasard parmi celles possibles
            x2, y2 = x + dx, y + dy
            id1 = self.cells[y][x]['id']
            id2 = self.cells[y2][x2]['id']
            if id1 != id2:  # Si les cellules ont des id différents, les fusionner
                self.ouvrir(x, y, x2, y2)
                old_id, new_id = max(id1, id2), min(id1, id2)
                for row in self.cells:  # Mettre à jour les id
                    for cell in row:
                        if cell['id'] == old_id:
                            cell['id'] = new_id
            if all(cell['id'] == ids[0] for row in self.cells for cell in row):
                break
        self.cells[0][0]['ouest'] = True  # Entrée
        self.cells[self.hauteur - 1][self.largeur - 1]['est'] = True  # Sortie

    def generer_exploration(self): # Algorithme d'exploration aléatoire
        stack = [(0, 0)] # Pile pour le backtracking
        visited = set(stack) # Ensemble des cellules visitées
        while stack: # Répéter jusqu'à ce que la pile soit vide
            x, y = stack[-1]
            directions = []
            if x > 0 and (x - 1, y) not in visited:
                directions.append((-1, 0))
            if x < self.largeur - 1 and (x + 1, y) not in visited:
                directions.append((1, 0))
            if y > 0 and (x, y - 1) not in visited:
                directions.append((0, -1))
            if y < self.hauteur - 1 and (x, y + 1) not in visited:
                directions.append((0, 1))
            if directions:
                dx, dy = choice(directions) # Choisir une direction au hasard parmi celles possibles
                x2, y2 = x + dx, y + dy
                self.ouvrir(x, y, x2, y2) # Ouvrir le mur entre les deux cellules
                stack.append((x2, y2)) # Ajouter la nouvelle cellule à la pile
                visited.add((x2, y2)) # Marquer la nouvelle cellule comme visitée
            else:
                stack.pop() # Toutes les cellules adjacentes ont été visitées, continuer le backtracking
        self.cells[0][0]['ouest'] = True  # Entrée
        self.cells[self.hauteur - 1][self.largeur - 1]['est'] = True  # Sortie

    def afficher(self, solved = False, terminal = False): # Affiche le labyrinthe dans le terminal ou renvoie un labyrinthe en texte
        texte = '' # Chaîne de caractères représentant le labyrinthe
        for y in range(self.hauteur):
            # Afficher le mur nord
            for x in range(self.largeur):
                texte += "+"
                if self.cells[y][x]['nord']:
                    if solved and (x, y) in self.chemin and (x, y - 1) in self.chemin:
                        texte += " ❚ "
                    else:
                        texte += "   "
                else:
                    texte += "---"
            texte += "+\n"
            # Afficher les murs ouest et est
            for x in range(self.largeur):
                if self.cells[y][x]['ouest']:
                    if solved and (x, y) in self.chemin and (x - 1, y) in self.chemin: # Mur ouvert et dans le chemin
                        texte += "━"
                    else:
                        texte += " "
                else:
                    texte += "|"
                if solved and (x, y) in self.chemin: # Cellule dans le chemin : afficher le caractère spécial
                    index = self.chemin.index((x, y))
                    previous = self.chemin[index - 1] if index > 0 else (-1, 0)
                    previous = (x - previous[0], y - previous[1])
                    next = self.chemin[index + 1] if index < len(self.chemin) - 1 else (self.largeur, self.hauteur - 1)
                    next = (x - next[0], y - next[1])
                    if previous in [(-1, 0), (1, 0)] and next in [(-1, 0), (1, 0)]:
                        texte += "━━━"
                    elif previous in [(0, -1), (0, 1)] and next in [(0, -1), (0, 1)]:
                        texte += " ❚ "
                    elif previous in [(-1, 0), (0, 1)] and next in [(-1, 0), (0, 1)]:
                        texte += " ┗━"
                    elif previous in [(1, 0), (0, 1)] and next in [(1, 0), (0, 1)]:
                        texte += "━┛ "
                    elif previous in [(-1, 0), (0, -1)] and next in [(-1, 0), (0, -1)]:
                        texte += " ┏━"
                    elif previous in [(1, 0), (0, -1)] and next in [(1, 0), (0, -1)]:
                        texte += "━┓ "
                else:
                    texte += "   "
            if self.cells[y][self.largeur - 1]['est']:
                if solved and (self.largeur - 1, y) in self.chemin and (self.largeur, y) in self.chemin:
                    texte += "━\n"
                else:
                    texte += " \n"
            else:
                texte += "|\n"
        # Afficher le mur sud
        for x in range(self.largeur):
            texte += "+"
            if self.cells[self.hauteur - 1][x]['sud']:
                texte += "   "
            else:
                texte += "---"
        texte += "+\n"

        if terminal:
            print(texte)
        return texte

    def solve(self): # Résout le labyrinthe en trouvant un chemin de l'entrée à la sortie
        start = (0,0)
        end = (self.largeur - 1, self.hauteur - 1)
        stack = [start]
        visited = set()
        parent = {start: None} # Pour reconstruire le chemin
        while stack:
            x, y = stack.pop()
            if (x, y) == end:
                break
            visited.add((x, y))
            directions = [] # Directions possibles
            if x > 0 and (x - 1, y) not in visited and self.cells[y][x]['ouest']:
                directions.append((-1, 0))
            if x < self.largeur - 1 and (x + 1, y) not in visited and self.cells[y][x]['est']:
                directions.append((1, 0))
            if y > 0 and (x, y - 1) not in visited and self.cells[y][x]['nord']:
                directions.append((0, -1))
            if y < self.hauteur - 1 and (x, y + 1) not in visited and self.cells[y][x]['sud']:
                directions.append((0, 1))
            for dx, dy in directions: # Explorer les directions possibles
                x2, y2 = x + dx, y + dy
                if (x2, y2) not in visited: # Si la cellule n'a pas été visitée
                    stack.append((x2, y2))
                    visited.add((x2, y2))
                    parent[(x2, y2)] = (x, y) # Enregistrer le parent

        # Reconstruire le chemin
        chemin = []
        while end:
            chemin.append(end)
            end = parent.get(end)
        chemin.reverse()
        chemin.insert(0, (-1, 0))
        chemin.append((self.largeur, self.hauteur - 1))
        self.chemin = chemin

    def generate_word_bytes(self, basic=True, with_solved=True):  
        def remove_border(cell, side): # Supprime la bordure spécifiée de la cellule
            """side: 'top','left','bottom','right' -> met val='nil' pour supprimer ce bord."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            el = OxmlElement('w:' + side)
            el.set(qn('w:val'), 'nil')
            tcBorders.append(el)

        document = Document()

        # orientation auto
        section = document.sections[-1]
        if self.largeur > self.hauteur:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        # marges fixes
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)

        # titre
        heading = document.add_heading(
            f'Labyrinthe {self.hauteur}×{self.largeur}'
            + (' avec solution' if with_solved and basic else ' résolu' if with_solved else ''),
            0
        )
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # calcul surface utile en cm
        usable_w_cm = (section.page_width.cm - section.left_margin.cm - section.right_margin.cm)
        usable_h_cm = (section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm)

        # taille maximale par cellule en cm
        cell_size_cm = min(usable_w_cm / self.largeur, usable_h_cm / self.hauteur)
        if cell_size_cm <= 0:
            cell_size_cm = 1.0

        def symbol_for_cell(x, y):
            chemin = getattr(self, 'chemin', [])
            if (x, y) not in chemin:
                return None
            i = chemin.index((x, y))
            prev = chemin[i-1] if i > 0 else None
            nxt = chemin[i+1] if i < len(chemin)-1 else None
            d1 = (prev[0] - x, prev[1] - y) if prev else None
            d2 = (nxt[0] - x, nxt[1] - y) if nxt else None
            if d1 is None and d2:
                return '━' if d2 in [(-1,0),(1,0)] else '┃'
            if d2 is None and d1:
                return '━' if d1 in [(-1,0),(1,0)] else '┃'
            s = {d1, d2}
            if s <= {(-1,0),(1,0)}: return '━'
            if s <= {(0,-1),(0,1)}: return '┃'
            if s == {(1,0),(0,1)}: return '┏'
            if s == {(-1,0),(0,1)}: return '┓'
            if s == {(1,0),(0,-1)}: return '┗'
            if s == {(-1,0),(0,-1)}: return '┛'
            return '╋'

        def draw_table(show_solution=False): # Dessine le labyrinthe dans un tableau Word
            table = document.add_table(rows=self.hauteur, cols=self.largeur) # Crée un tableau
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # colonnes
            for col in table.columns:
                col.width = Cm(cell_size_cm)

            # lignes / cellules
            for r_idx, row in enumerate(table.rows):
                try:
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    row.height = Cm(cell_size_cm)
                except Exception:
                    pass  # row.height peut bugger ; Word le gèrera
                for c_idx, cell in enumerate(row.cells):
                    # supprimer les bordures correspondant aux ouvertures
                    if self.cells[r_idx][c_idx]['nord']:
                        remove_border(cell, 'top')
                    if self.cells[r_idx][c_idx]['sud']:
                        remove_border(cell, 'bottom')
                    if self.cells[r_idx][c_idx]['ouest']:
                        remove_border(cell, 'left')
                    if self.cells[r_idx][c_idx]['est']:
                        remove_border(cell, 'right')

                    # centrage
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # solution avec box-drawing
                    if show_solution:
                        sym = symbol_for_cell(c_idx, r_idx)
                        if sym:
                            p = cell.paragraphs[0]
                            run = p.add_run(sym)
                            run.bold = True
                            # taille en points basée sur la taille de la cellule (1 cm ≈ 28.35 pt)
                            run.font.size = Pt(max(8, int(cell_size_cm * 28)))
                            run.font.name = "Courier New"

        if basic:
            draw_table(show_solution=False)

        if with_solved:
            if basic:
                document.add_page_break()
            draw_table(show_solution=True)

        # Sauvegarde en mémoire
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    def generate_pdf_bytes(self, basic=True, with_solved=True):
        # dimensions PDF
        page_width, page_height = landscape(A4)

        # marges
        margin = 1*cm
        usable_w = page_width - 2*margin
        usable_h = page_height - 2*margin

        # taille cellule
        cell_size = min((usable_w-100) / self.largeur, (usable_h-100) / self.hauteur)

        # Créer un buffer en mémoire
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=landscape(A4))
        c.setFont("Helvetica", 12)

        c.setTitle(f'Labyrinthe {self.hauteur}x{self.largeur}' + (' avec solution' if with_solved and basic else ' résolu' if with_solved else ''))
        c.setAuthor("Générateur de Labyrinthe Python - Xavier Labarre - github.com/labarrex")
        c.setSubject(f"Labyrinthe {self.largeur}x{self.hauteur} généré aléatoirement")

        def draw_maze(show_solution=False):
            # titre
            titre = f"Labyrinthe {self.hauteur}×{self.largeur}"
            if with_solved and basic and show_solution:
                titre += " avec solution"
            elif with_solved and show_solution:
                titre += " résolu"

            # titre en haut de page
            c.setFont("Helvetica-Bold", 30)
            c.drawCentredString(page_width/2, page_height - 2.3*cm, titre)

            # zone labyrinthe
            lab_w = self.largeur * cell_size
            lab_h = self.hauteur * cell_size
            offset_x = (page_width - lab_w) / 2
            offset_y = (page_height - lab_h) / 2 - 1.1*cm

            # murs
            for y in range(self.hauteur):
                for x in range(self.largeur):
                    cell = self.cells[y][x]
                    px = offset_x + x*cell_size
                    py = offset_y + (self.hauteur-1-y)*cell_size

                    if not cell['nord']:
                        c.line(px, py+cell_size, px+cell_size, py+cell_size)
                    if not cell['sud']:
                        c.line(px, py, px+cell_size, py)
                    if not cell['ouest']:
                        c.line(px, py, px, py+cell_size)
                    if not cell['est']:
                        c.line(px+cell_size, py, px+cell_size, py+cell_size)

            # solution en chemin vectoriel
            if show_solution and hasattr(self, 'chemin') and self.chemin:
                path = c.beginPath()
                first = True
                for (x, y) in self.chemin:
                    if x < 0 or y < 0 or x >= self.largeur or y >= self.hauteur:
                        continue
                    cx = offset_x + x*cell_size + cell_size/2
                    cy = offset_y + (self.hauteur-1-y)*cell_size + cell_size/2
                    if first:
                        path.moveTo(cx, cy)
                        first = False
                    else:
                        path.lineTo(cx, cy)
                if not first:
                    c.setStrokeColorRGB(1, 0, 0)
                    c.setLineWidth(2)
                    c.drawPath(path)

        # version simple
        if basic:
            draw_maze(show_solution=False)

        # version avec solution
        if with_solved:
            if basic:
                c.showPage()
            draw_maze(show_solution=True)

        c.save()
        buffer.seek(0)
        return buffer.getvalue()


    def generate_svg(self, solved=False, width=600, height=600): # Génère un SVG du labyrinthe
        cell_size = min(width / self.largeur, height / self.hauteur)
        maze_width = self.largeur * cell_size
        maze_height = self.hauteur * cell_size
        
        # Centrer le labyrinthe
        offset_x = (width - maze_width) / 2 + 10
        offset_y = (height - maze_height) / 2 + 10

        svg_lines = []
        svg_lines.append(f'<svg width="{width+20}" height="{height+20}" viewBox="0 0 {width+20} {height+20}" xmlns="http://www.w3.org/2000/svg">')
        svg_lines.append('<style>')
        svg_lines.append('.maze-wall { stroke: #2c3e50; stroke-width: 1; }')
        svg_lines.append('.maze-path { stroke: #e74c3c; stroke-width: 3; stroke-linecap: round; fill: none; }')
        svg_lines.append('.maze-bg { fill: #ffffff; }')
        svg_lines.append('</style>')
        
        # Fond blanc
        svg_lines.append(f'<rect x="{offset_x}" y="{offset_y}" width="{maze_width}" height="{maze_height}" class="maze-bg"/>')
        
        # Dessiner les murs
        for y in range(self.hauteur):
            for x in range(self.largeur):
                cell = self.cells[y][x]
                px = offset_x + x * cell_size
                py = offset_y + y * cell_size

                # Mur nord
                if not cell['nord']:
                    svg_lines.append(f'<line x1="{px}" y1="{py}" x2="{px + cell_size}" y2="{py}" class="maze-wall"/>')
                
                # Mur sud
                if not cell['sud']:
                    svg_lines.append(f'<line x1="{px}" y1="{py + cell_size}" x2="{px + cell_size}" y2="{py + cell_size}" class="maze-wall"/>')
                
                # Mur ouest
                if not cell['ouest']:
                    svg_lines.append(f'<line x1="{px}" y1="{py}" x2="{px}" y2="{py + cell_size}" class="maze-wall"/>')
                
                # Mur est
                if not cell['est']:
                    svg_lines.append(f'<line x1="{px + cell_size}" y1="{py}" x2="{px + cell_size}" y2="{py + cell_size}" class="maze-wall"/>')
        
        # Dessiner la solution si demandée
        if solved and hasattr(self, 'chemin'):
            path_coords = []
            for i, (x, y) in enumerate(self.chemin):
                # Ignorer les points virtuels d'entrée/sortie
                if x == -1 or x == self.largeur or y == -1 or y == self.hauteur:
                    continue

                px = offset_x + x * cell_size + cell_size/2
                py = offset_y + y * cell_size + cell_size/2
                path_coords.append(f"{px},{py}")
            
            if path_coords:
                path_string = "M " + " L ".join(path_coords)
                svg_lines.append(f'<path d="{path_string}" class="maze-path"/>')
        
        svg_lines.append('</svg>')
        return '\n'.join(svg_lines)

    def generate_svg_bytes(self, solved=False, width=None, height=None):
        # Générer le SVG en texte HTML
        svg_html = self.generate_svg(solved=solved, width=width, height=height)
        
        # Convertir le texte en bytes (encodage UTF-8)
        svg_bytes = svg_html.encode('utf-8')
        
        return svg_bytes
